"""Exportación de reportes a un documento de Word (.docx).

Se genera en memoria y se devuelve como bytes: los documentos son pequeños
(decenas de reportes) y así no hay archivos temporales que limpiar ni estado
compartido entre peticiones.

Cada reporte se imprime como un CUADRO DE FICHA —una tabla de Word— y no como
párrafos sueltos. La tabla es lo único que Word, Pages y Google Docs pintan
igual: el relleno de la banda, las reglas de la rejilla y, sobre todo,
`cantSplit`, que impide que una ficha se parta entre dos hojas. Debajo del
cuadro va el cuerpo de la nota en párrafos normales.

La especificación de estilos vive en `docs/EXPORT_DOCX.md`, junto a la
plantilla de referencia.
"""
from __future__ import annotations

import io
import re
from datetime import date
from importlib.resources import files

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.shared import Pt, RGBColor, Twips
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from odin.analysis.text_norm import norm_key
from odin.api import deps
from odin.db.models import Article
from odin.scrapers import source_name

# Etiquetas legibles para los códigos que guarda el análisis. El documento lo
# lee una persona, no un programa: "Negativo" dice más que "NEG".
_SENTIMENT_LABELS = {"POS": "Positivo", "NEG": "Negativo", "NEU": "Neutro"}

#: Lo que se imprime donde no hay dato. Nunca cadena vacía ni "None": una celda
#: en blanco se lee como un error de la exportación, un guion se lee como
#: "no había".
_VACIO = "—"


# --- Paleta y métricas del cuadro -------------------------------------------
# Los colores no son decorativos: el único color fuerte es la banda del
# encabezado, y todo lo descriptivo queda en grises. Un reporte de prensa que
# usa color en los datos sugiere una jerarquía que los datos no tienen.

_BANDA = "22262A"        # relleno de la banda superior
_ENTIDADES = "F4F5F6"    # relleno de la franja de entidades
_REGLA = "C9CED3"        # 1 pt: marco del cuadro y separadores mayores
_REGLA_SUAVE = "E1E4E7"  # 0.5 pt: rejilla interna de metadatos
_REGLA_TENUE = "DCDFE2"  # 0.5 pt: separador de la fuente
_TINTA = "22262A"        # valores de metadatos y entidades
_TINTA_TITULAR = "141719"
_ETIQUETA = "7A8086"     # etiquetas de la rejilla
_APAGADO = "9AA0A6"      # etiqueta "FUENTE"
_BLANCO = "FFFFFF"

_UI = "Helvetica Neue"   # metadatos y etiquetas (fallback Arial en fontTable)
_SERIF = "Georgia"       # titular y cuerpo

#: Ancho de columna en twips. Dos columnas = 9360 = 6.5" = el ancho útil entre
#: los márgenes de una hoja Letter con 1" a cada lado.
_COLUMNA = 4680
_ANCHO = _COLUMNA * 2


# --- Utilidades de OOXML ----------------------------------------------------
# python-docx no expone relleno de celda, bordes por celda, `cantSplit` ni
# tracking de caracteres, así que esos cuatro se escriben a mano. Word valida
# el ORDEN de los hijos de cada elemento de propiedades: un `w:tcBorders`
# colocado después de un `w:tcMar` hace que el documento abra con advertencia,
# que es justamente lo que el criterio de aceptación prohíbe. De ahí que cada
# inserción declare ante qué hermanos debe quedar.


def _tag(name: str, **attrs: str):
    """Un elemento `w:<name>` con atributos del mismo namespace."""
    element = OxmlElement(f"w:{name}")
    for key, value in attrs.items():
        element.set(qn(f"w:{key}"), value)
    return element


def _insert(parent, child, before: tuple[str, ...]) -> None:
    """Coloca `child` antes del primero de `before` que exista, o al final."""
    for name in before:
        hermano = parent.find(qn(f"w:{name}"))
        if hermano is not None:
            hermano.addprevious(child)
            return
    parent.append(child)


#: Los bordes se declaran como datos `(lado, grosor en pt, color)` y no como
#: elementos ya construidos: lxml MUEVE un elemento al añadirlo a un segundo
#: padre, así que una misma regla reusada en dos celdas dejaría a la primera sin
#: línea. Como datos, cada celda construye la suya.
_Borde = tuple[str, float, str]

_ABAJO_FUERTE: _Borde = ("bottom", 1, _REGLA)
_ABAJO_SUAVE: _Borde = ("bottom", 0.5, _REGLA_SUAVE)
_IZQUIERDA_SUAVE: _Borde = ("left", 0.5, _REGLA_SUAVE)
_ARRIBA_FUERTE: _Borde = ("top", 1, _REGLA)
_ARRIBA_TENUE: _Borde = ("top", 0.5, _REGLA_TENUE)

#: Word exige este orden entre los hijos de `w:tcBorders` / `w:tblBorders`.
_LADOS = ("top", "left", "bottom", "right", "insideH", "insideV")


def _bordes(name: str, especificaciones: tuple[_Borde, ...]):
    """Un `w:tcBorders`/`w:tblBorders` con los lados en el orden del esquema."""
    contenedor = _tag(name)
    for lado, pt, color in sorted(especificaciones, key=lambda b: _LADOS.index(b[0])):
        # `w:sz` va en octavos de punto; "none" es un lado explícitamente sin
        # línea, distinto de omitirlo (que heredaría el borde de la tabla).
        if pt:
            contenedor.append(
                _tag(lado, val="single", sz=str(round(pt * 8)), space="0", color=color)
            )
        else:
            contenedor.append(_tag(lado, val="none", sz="0", space="0", color="auto"))
    return contenedor


# Hermanos que siguen a cada elemento dentro de su contenedor, según el esquema
# de WordprocessingML. Recortados a lo que aparece en este documento.
_TRAS_TCBORDERS = ("shd", "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark")
_TRAS_SHD = ("noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark")
_TRAS_TCMAR = ("textDirection", "tcFitText", "vAlign", "hideMark")
_TRAS_TBLBORDERS = ("shd", "tblLayout", "tblCellMar", "tblLook")
_TRAS_TBLCELLMAR = ("tblLook",)
_TRAS_CANTSPLIT = ("trHeight", "tblHeader", "tblCellSpacing", "jc", "hidden", "ins", "del")
_TRAS_RSPACING = (
    "w", "kern", "position", "sz", "szCs", "highlight", "u", "effect", "bdr",
    "shd", "fitText", "vertAlign", "rtl", "cs", "em", "lang",
)


def _run(
    paragraph: Paragraph,
    text: str,
    *,
    font: str,
    size: float,
    color: str,
    bold: bool = False,
    tracking: int = 0,
):
    """Un fragmento de texto con formato directo.

    TODO el texto del documento pasa por aquí, también el de la portada y el
    cuerpo, que además llevan un estilo de párrafo. Es deliberado: el lector de
    .docx de macOS ignora los estilos de párrafo con nombre propio —el título
    sale en Times 12— y solo respeta el formato directo. El estilo queda como
    portador del espaciado; la tipografía la fija esta función, y así se ve
    igual en las tres aplicaciones.

    `tracking` va en veinteavos de punto, la unidad de `w:spacing`.
    """
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    # Solo se marca la negrita cuando la hay: apagarla explícitamente escribe
    # `<w:b w:val="0"/>`, y el lector de macOS lo interpreta como negrita
    # ENCENDIDA, con lo que el documento entero saldría en bold.
    if bold:
        run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    if tracking:
        _insert(run._element.get_or_add_rPr(), _tag("spacing", val=str(tracking)), _TRAS_RSPACING)
    return run


def _parrafo(cell: _Cell, *, primero: bool, line: float | None = None) -> Paragraph:
    """Un párrafo dentro de una celda, sin espaciado heredado.

    La celda nace con un párrafo vacío; el primer texto lo reusa y el resto se
    añade, para no dejar una línea en blanco arriba.
    """
    paragraph = cell.paragraphs[0] if primero else cell.add_paragraph()
    formato = paragraph.paragraph_format
    formato.space_before = Pt(0)
    formato.space_after = Pt(0)
    if line:
        formato.line_spacing = line
    return paragraph


def _celda(
    cell: _Cell,
    *,
    ancho: int = _COLUMNA,
    relleno: str | None = None,
    bordes: tuple[_Borde, ...] = (),
    arriba: int,
    abajo: int,
) -> None:
    """Ancho, relleno, bordes y padding vertical de una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.get_or_add_tcW().set(qn("w:w"), str(ancho))
    tcPr.get_or_add_tcW().set(qn("w:type"), "dxa")

    if bordes:
        _insert(tcPr, _bordes("tcBorders", bordes), _TRAS_TCBORDERS)

    if relleno:
        _insert(tcPr, _tag("shd", val="clear", color="auto", fill=relleno), _TRAS_SHD)

    margenes = _tag("tcMar")
    margenes.append(_tag("top", w=str(arriba), type="dxa"))
    margenes.append(_tag("bottom", w=str(abajo), type="dxa"))
    _insert(tcPr, margenes, _TRAS_TCMAR)


def _fila(table: Table, *, combinada: bool = False) -> list[_Cell]:
    """Una fila que no se parte entre hojas.

    `cantSplit` por fila y no `keepNext` por párrafo: es lo que las tres apps
    respetan igual, y es lo que evita que una ficha quede cortada al pie.
    """
    row = table.add_row()
    _insert(row._tr.get_or_add_trPr(), _tag("cantSplit"), _TRAS_CANTSPLIT)
    if combinada:
        return [row.cells[0].merge(row.cells[1])]
    return list(row.cells)


def _marco(table: Table) -> None:
    """Marco exterior del cuadro y márgenes de celda por defecto.

    Las reglas internas se declaran celda por celda (`insideH`/`insideV` en
    `none`): la rejilla de metadatos no lleva línea en todos los cruces, y una
    regla horizontal global pintaría también las que sobran.
    """
    table.autofit = False
    tblPr = table._tbl.tblPr

    marco = tuple((lado, 1, _REGLA) for lado in ("top", "left", "bottom", "right"))
    dentro = tuple((lado, 0, "auto") for lado in ("insideH", "insideV"))
    _insert(tblPr, _bordes("tblBorders", marco + dentro), _TRAS_TBLBORDERS)

    margenes = _tag("tblCellMar")
    margenes.append(_tag("top", w="0", type="dxa"))
    margenes.append(_tag("left", w="150", type="dxa"))
    margenes.append(_tag("bottom", w="0", type="dxa"))
    margenes.append(_tag("right", w="150", type="dxa"))
    _insert(tblPr, margenes, _TRAS_TBLCELLMAR)

    for columna in table.columns:
        columna.width = Twips(_COLUMNA)


# --- Limpieza del cuerpo ----------------------------------------------------

# Un ladillo del medio ("Alegada negligencia", "POLÍTICA") es corto y no cierra
# en punto. El umbral es generoso a propósito: descartar prosa real por error es
# peor que dejar pasar un ladillo, porque el cliente no puede recuperar lo que
# no se exportó.
_KICKER_MAX_CHARS = 60

#: Cualquier blanco, incluidos el no separable y el fino que trae el scrape:
#: `\s` sobre `str` cubre todo el espacio en blanco Unicode, no solo el ASCII.
_ESPACIOS = re.compile(r"\s+")

#: Una comilla recta abre si viene tras espacio o apertura, y cierra si no. Se
#: mira el carácter anterior en vez de alternar a ciegas porque un párrafo
#: puede llegar con una comilla desemparejada y a partir de ahí quedarían todas
#: al revés.
_COMILLA = re.compile(r'(^|[\s(\[¡¿«—–-])"|"')


def _is_section_kicker(paragraph: str) -> bool:
    """¿Es un ladillo de sección y no un párrafo de la nota?

    Se exige que sea corto Y que no termine en puntuación de cierre: una frase
    breve pero terminada ("«El peso pudo más que el techo», afirmó.") es prosa
    y se conserva.
    """
    if len(paragraph) > _KICKER_MAX_CHARS:
        return False
    return not paragraph.rstrip().endswith((".", "?", "!", "…", '"', "”", "»"))


def normalize(paragraph: str) -> str:
    """Tipografía del párrafo, antes de imprimirlo.

    En pantalla un espacio doble o una comilla recta pasan; impresos en un
    documento que el cliente entrega, se ven como descuido. Se colapsa todo
    espacio en blanco a uno solo y las comillas rectas se vuelven curvas.
    """
    texto = _ESPACIOS.sub(" ", paragraph).strip()
    return _COMILLA.sub(lambda m: f"{m.group(1)}“" if m.group(1) else "”", texto)


def clean_body(body: str | None, *, title: str, topic: str | None) -> list[str]:
    """Párrafos del cuerpo, sin lo que arrastra el scrape.

    Lo que viene del sitio trae, antes de la nota, el tema suelto y el titular
    repetido; a veces el sumario aparece dos veces; y en el medio quedan
    ladillos de sección. En pantalla se toleran, pero el .docx lo imprime y
    entrega el cliente, y ahí se leen como errores.

    Reglas, las del README de la plantilla:
      - fuera el tema si aparece como párrafo
      - fuera el titular repetido
      - de un párrafo repetido queda solo la ÚLTIMA aparición
      - fuera los ladillos de sección
      - comillas y espacios normalizados

    Se conserva la última y no la primera por el caso que la motiva: el sumario
    viene suelto ANTES de la nota y otra vez en su lugar narrativo. Quedándose
    con la primera, el cuerpo abre resumiendo lo que va a contar; con la última,
    arranca donde arranca la nota.

    La comparación es insensible a mayúsculas y acentos porque el mismo texto
    vuelve del scrape con otra caja más seguido de lo que parece.
    """
    if not body:
        return []

    descartar = {norm_key(t) for t in (title, topic) if t}

    candidatos = [
        (norm_key(p), p)
        for p in (normalize(raw) for raw in body.split("\n"))
        if p and norm_key(p) not in descartar and not _is_section_kicker(p)
    ]

    # Recorrido inverso para quedarse con la ÚLTIMA aparición de cada repetido,
    # y luego se restaura el orden original.
    vistos: set[str] = set()
    salida: list[str] = []
    for key, paragraph in reversed(candidatos):
        if key in vistos:
            continue
        vistos.add(key)
        salida.append(paragraph)

    return list(reversed(salida))


# --- Datos ------------------------------------------------------------------

_MESES = (
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
)


def _format_date(value: date | None) -> str:
    """Día/mes/año. Sin hora: es lo que pidió el cliente y lo que la columna
    `analyzed_on` guarda."""
    return value.strftime("%d/%m/%Y") if value else _VACIO


def _long_date(value: date) -> str:
    return f"{value.day} de {_MESES[value.month - 1]} de {value.year}"


def _period(articles: list[Article]) -> str:
    """Rango de publicación de los reportes, en prosa.

    Se acorta lo que se repite —"21–28 de agosto de 2026" y no "21 de agosto de
    2026 – 28 de agosto de 2026"— porque va en el subtítulo de la portada, en
    una sola línea junto al conteo y los medios.
    """
    fechas = sorted(a.published_at.date() for a in articles if a.published_at)
    if not fechas:
        return ""

    inicio, fin = fechas[0], fechas[-1]
    if inicio == fin:
        return _long_date(inicio)
    if (inicio.year, inicio.month) == (fin.year, fin.month):
        return f"{inicio.day}–{fin.day} de {_MESES[fin.month - 1]} de {fin.year}"
    if inicio.year == fin.year:
        return f"{inicio.day} de {_MESES[inicio.month - 1]} – {_long_date(fin)}"
    return f"{_long_date(inicio)} – {_long_date(fin)}"


def _media(articles: list[Article]) -> str:
    """Medios distintos, en el orden en que aparecen en la selección."""
    nombres: list[str] = []
    for article in articles:
        nombre = source_name(article.source)
        if nombre and nombre not in nombres:
            nombres.append(nombre)
    return ", ".join(nombres)


def _entities(article: Article) -> list[str]:
    """Nombres únicos, en el orden en que los devolvió el análisis.

    Sin ordenar alfabéticamente: el análisis las devuelve por relevancia y esa
    es la lectura útil de un vistazo.
    """
    vistos: set[str] = set()
    nombres: list[str] = []
    for entity in sorted(article.entities, key=lambda e: e.id or 0):
        clave = norm_key(entity.name)
        if clave in vistos:
            continue
        vistos.add(clave)
        nombres.append(entity.name)
    return nombres


def _value(text: str | None) -> str:
    return text.strip() if text and text.strip() else _VACIO


# --- Composición ------------------------------------------------------------

#: Plantilla con el tamaño de hoja, los márgenes, el encabezado corrido y el
#: pie con número de página. Se usa como archivo BASE en vez de recrear eso en
#: código: es la parte del documento que se ajusta editando el .docx y no
#: tocando Python. La tipografía, en cambio, la fija el código —ver `_run`—.
#: Empaquetada vía `[tool.setuptools.package-data]` en pyproject.toml.
_TEMPLATE = files("odin.exports") / "reportes-odin-template.docx"


# `docx.Document` (importado arriba) es la función fábrica que crea el
# documento; la clase real, la que hace falta para anotar tipos, vive en
# `docx.document.Document` — de ahí el alias `DocxDocument`. Sin él mypy
# rechaza la anotación porque una función no es un tipo válido.
def _load_template() -> DocxDocument:
    """Abre la plantilla y le quita los reportes de ejemplo.

    Se borra el cuerpo pero NO la `sectPr` final, que es donde viven el tamaño
    de hoja, los márgenes y las referencias al encabezado y al pie: quitarla
    devolvería el documento a los valores por defecto de Word.
    """
    with _TEMPLATE.open("rb") as handle:
        document = Document(handle)

    body = document.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)
    return document


def _cover(document: DocxDocument, articles: list[Article]) -> None:
    """Portada: qué es el documento y qué abarca, en dos líneas."""
    titulo = document.add_paragraph(style="ODIN Title")
    _run(titulo, "Reportes de prensa", font=_SERIF, size=24, color="111111", bold=True)

    total = f"{len(articles)} {'reporte' if len(articles) == 1 else 'reportes'}"
    partes = ["ODIN", _period(articles), total, _media(articles)]
    subtitulo = document.add_paragraph(style="ODIN Subtitle")
    _run(subtitulo, " · ".join(p for p in partes if p), font=_UI, size=9, color="666666")


def _band(cells: list[_Cell], number: int, article: Article) -> None:
    """Fila 1: número del reporte y su sentimiento, en blanco sobre oscuro."""
    izquierda, derecha = cells

    _celda(izquierda, relleno=_BANDA, bordes=(_ABAJO_FUERTE,), arriba=90, abajo=90)
    _run(
        _parrafo(izquierda, primero=True),
        f"REPORTE {number:02d}",
        font=_UI, size=8, color=_BLANCO, bold=True, tracking=24,
    )

    _celda(derecha, relleno=_BANDA, bordes=(_ABAJO_FUERTE,), arriba=90, abajo=90)
    parrafo = _parrafo(derecha, primero=True)
    parrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    etiqueta = _SENTIMENT_LABELS.get(article.overall_sentiment or "", _VACIO)
    _run(
        parrafo,
        etiqueta.upper(),
        font=_UI, size=7.5, color=_BLANCO, bold=True, tracking=28,
    )


def _headline(cell: _Cell, article: Article) -> None:
    """Fila 2: el titular, lo único del cuadro que se lee de lejos."""
    _celda(
        cell,
        ancho=_ANCHO,
        bordes=(_ABAJO_FUERTE,),
        arriba=170,
        abajo=150,
    )
    _run(
        _parrafo(cell, primero=True, line=1.2),
        _value(article.title),
        font=_SERIF, size=16, color=_TINTA_TITULAR, bold=True,
    )


def _field(cell: _Cell, label: str, value: str, *, bordes: tuple[_Borde, ...]) -> None:
    """Una casilla de la rejilla: etiqueta menuda arriba, dato debajo."""
    _celda(cell, bordes=bordes, arriba=120, abajo=120)

    etiqueta = _parrafo(cell, primero=True)
    etiqueta.paragraph_format.space_after = Pt(1)
    _run(etiqueta, label.upper(), font=_UI, size=7, color=_ETIQUETA, bold=True, tracking=22)

    # "Semibold" no existe como peso en OOXML —solo hay negrita o no—, así que
    # el dato va en negrita: a 9.5 pt junto a una etiqueta de 7 pt en gris, es
    # lo que produce el contraste que pide el diseño.
    _run(_parrafo(cell, primero=False), value, font=_UI, size=9.5, color=_TINTA, bold=True)


def _meta_grid(table: Table, article: Article) -> None:
    """Filas 3–5: los seis datos del reporte, en dos columnas.

    El orden es fijo y se lee en Z (izquierda a derecha, arriba a abajo): quien
    revisa treinta fichas seguidas busca el medio siempre en el mismo sitio.
    """
    publicado = article.published_at.date() if article.published_at else None
    datos = [
        ("Medio", _value(source_name(article.source))),
        ("Sección", _value(article.section)),
        ("Publicado", _format_date(publicado)),
        ("Tema", _value(article.main_topic)),
        (
            "Documentalista",
            article.documentalist.display_name if article.documentalist else "Automático",
        ),
        ("Analizado", _format_date(article.analyzed_on)),
    ]

    for indice in range(0, len(datos), 2):
        izquierda, derecha = _fila(table)
        # La última fila no lleva regla abajo: la cierra el borde del cuadro o
        # la franja de entidades, y doblarla se vería como una línea gruesa.
        ultima = indice + 2 >= len(datos)
        regla: tuple[_Borde, ...] = () if ultima else (_ABAJO_SUAVE,)

        _field(izquierda, *datos[indice], bordes=regla)
        # La regla vertical vive en la celda de la derecha: un borde izquierdo
        # ahí y no uno derecho aquí, para que no se dupliquen en el mismo cruce.
        _field(derecha, *datos[indice + 1], bordes=(*regla, _IZQUIERDA_SUAVE))


def _entities_row(cell: _Cell, nombres: list[str]) -> None:
    """Fila 6: de quién habla la nota. Solo si el análisis encontró alguien."""
    _celda(
        cell,
        ancho=_ANCHO,
        relleno=_ENTIDADES,
        bordes=(_ARRIBA_FUERTE,),
        arriba=140,
        abajo=140,
    )

    etiqueta = _parrafo(cell, primero=True)
    etiqueta.paragraph_format.space_after = Pt(1)
    _run(
        etiqueta,
        f"ENTIDADES MENCIONADAS · {len(nombres)}",
        font=_UI, size=7, color=_ETIQUETA, bold=True, tracking=22,
    )
    # Punto medio y no coma: hay entidades que llevan coma adentro
    # ("Corporación del Acueducto y Alcantarillado, Santiago").
    _run(
        _parrafo(cell, primero=False, line=1.4),
        " · ".join(nombres),
        font=_UI, size=9, color=_TINTA,
    )


def _source_row(cell: _Cell, article: Article) -> None:
    """Fila 7: la URL, para que el reporte impreso siga siendo verificable."""
    _celda(
        cell,
        ancho=_ANCHO,
        bordes=(_ARRIBA_TENUE,),
        arriba=110,
        abajo=110,
    )
    parrafo = _parrafo(cell, primero=True)
    _run(parrafo, "FUENTE", font=_UI, size=7, color=_APAGADO, bold=True, tracking=22)
    # Sin hipervínculo: el azul subrayado de Word sería el único color de la
    # ficha fuera de la banda, y en papel no sirve de nada.
    _run(parrafo, f"   {_value(article.url)}", font=_UI, size=7.5, color=_ETIQUETA)


def _ficha(document: DocxDocument, article: Article, number: int) -> None:
    """El cuadro completo de un reporte."""
    table = document.add_table(rows=0, cols=2)
    _marco(table)

    _band(_fila(table), number, article)
    _headline(_fila(table, combinada=True)[0], article)
    _meta_grid(table, article)

    nombres = _entities(article)
    if nombres:
        _entities_row(_fila(table, combinada=True)[0], nombres)
    _source_row(_fila(table, combinada=True)[0], article)


def _body(document: DocxDocument, article: Article) -> None:
    """El texto de la nota, debajo del cuadro."""
    parrafos = clean_body(article.body, title=article.title or "", topic=article.main_topic)

    # `or [""]` para que siempre quede un párrafo detrás del cuadro: una tabla
    # pegada al final del documento —o a otra tabla— es lo que hace que Word
    # abra el archivo con advertencia.
    for posicion, texto in enumerate(parrafos or [""]):
        parrafo = document.add_paragraph(style="ODIN Body")
        formato = parrafo.paragraph_format
        formato.line_spacing = 1.5
        formato.space_after = Pt(10)
        if posicion == 0:
            formato.space_before = Pt(13)
        _run(parrafo, texto, font=_SERIF, size=11, color="1A1A1A")


def build_document(articles: list[Article]) -> bytes:
    """Arma el .docx con los reportes recibidos, en ese orden."""
    document = _load_template()
    _cover(document, articles)

    for index, article in enumerate(articles):
        # Salto explícito y no `page_break_before` en el párrafo del cuadro:
        # un cuadro es una tabla, y una tabla no admite esa propiedad.
        if index:
            document.add_page_break()
        _ficha(document, article, index + 1)
        _body(document, article)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_articles(article_ids: list[int]) -> bytes:
    """Busca los reportes pedidos y devuelve el documento.

    Los ids que ya no existen se ignoran en vez de fallar: entre que el usuario
    vio la lista y pulsó exportar, alguien pudo borrar uno, y perder la descarga
    entera por eso sería peor que entregar el resto. Solo se devuelve 404 si no
    queda ninguno.
    """
    if not article_ids:
        raise HTTPException(status_code=422, detail="No hay reportes seleccionados.")

    session = deps.get_session()
    try:
        rows = session.scalars(
            select(Article)
            .options(selectinload(Article.entities), selectinload(Article.documentalist))
            .where(Article.id.in_(article_ids))
        ).all()
        if not rows:
            raise HTTPException(
                status_code=404, detail="Ninguno de los reportes seleccionados existe."
            )

        # Respetar el orden en que llegaron los ids: es el que el usuario ve en
        # pantalla, y `IN (...)` no garantiza ninguno.
        by_id = {row.id: row for row in rows}
        ordered = [by_id[i] for i in article_ids if i in by_id]
        return build_document(ordered)
    finally:
        session.close()
