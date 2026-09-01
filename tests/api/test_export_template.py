"""El .docx exportado sigue la especificación de docs/EXPORT_DOCX.md.

Cada reporte se imprime como un CUADRO DE FICHA —una tabla de Word— y el cuerpo
de la nota va debajo, en párrafos. Estas pruebas miran la ficha por dentro: qué
va en cada fila, qué se pinta y qué NO se pinta, porque el documento se imprime
y se entrega al cliente y ahí un desajuste no se puede corregir después.

La plantilla se usa como archivo base para heredar el tamaño de hoja, los
márgenes, el encabezado corrido y el pie con número de página. La tipografía, en
cambio, la escribe el código en formato directo, porque hay lectores de .docx
que ignoran los estilos de párrafo con nombre propio.
"""
from __future__ import annotations

import io
from datetime import UTC, date, datetime

import pytest
from docx import Document
from docx.oxml.ns import qn

from odin.core import auth
from odin.core.auth import create_token
from odin.db.models import Article, Entity, User


def _headers():
    token, _ = create_token("jperez")
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def reports(sqlite_sessionmaker):
    """Dos reportes: uno completo y otro pelado.

    El segundo es el caso que rompe el diseño si nadie lo cuida: sin entidades
    y sin documentalista, la ficha no puede quedar con una fila vacía ni con un
    `None` impreso.
    """
    session = sqlite_sessionmaker()
    session.add(
        User(
            username="jperez",
            display_name="Juan Pérez",
            first_name="Juan",
            last_name="Pérez",
            password_hash=auth.hash_password("x", iterations=1000),
            role="documentalista",
        )
    )
    session.commit()
    autor = session.query(User).filter_by(username="jperez").one().id

    completo = Article(
        source="listin_diario",
        url="https://listindiario.com/n1",
        title="Titular del reporte 1",
        section="Provincias",
        body=(
            "tema suelto 1\n"
            "Titular del reporte 1\n"
            "Primer párrafo real de la nota, con su punto final.\n"
            "Alegada negligencia\n"
            "Segundo párrafo real de la nota, también terminado."
        ),
        main_topic="tema suelto 1",
        overall_sentiment="NEG",
        published_at=datetime(2026, 8, 21, tzinfo=UTC),
        documentalist_id=autor,
        analyzed_on=date(2026, 8, 28),
    )
    completo.entities = [
        Entity(name="Coraasan", type="ORG"),
        Entity(name="Ramón Peña", type="PERSON"),
    ]

    pelado = Article(
        source="diario_libre",
        url="https://diariolibre.com/n2",
        title="Titular del reporte 2",
        body="Párrafo único de la segunda nota.",
        overall_sentiment="POS",
        published_at=datetime(2026, 8, 28, tzinfo=UTC),
    )

    ids = []
    for article in (completo, pelado):
        session.add(article)
        session.commit()
        ids.append(article.id)
    session.close()
    return ids


def _export(api_client, ids) -> Document:
    resp = api_client.post("/api/articles/export", json={"article_ids": ids}, headers=_headers())
    assert resp.status_code == 200, resp.text
    return Document(io.BytesIO(resp.content))


def _rows(cuadro) -> list[list[str]]:
    """Las filas de una ficha como texto, sin repetir las celdas combinadas."""
    filas = []
    for row in cuadro.rows:
        textos = [c.text for c in row.cells]
        filas.append(textos[:1] if len(set(textos)) == 1 else textos)
    return filas


def _field(cell) -> tuple[str, str]:
    """Etiqueta y dato de una casilla de la rejilla."""
    return cell.paragraphs[0].text, cell.paragraphs[1].text


def _todo_el_texto(doc: Document) -> str:
    parrafos = [p.text for p in doc.paragraphs]
    celdas = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    return "\n".join(parrafos + celdas)


class TestTemplate:
    def test_keeps_the_running_header_and_page_number(self, api_client, reports):
        doc = _export(api_client, reports)
        section = doc.sections[0]

        assert "ODIN" in " ".join(p.text for p in section.header.paragraphs)
        assert "PAGE" in section.footer.paragraphs[0]._p.xml

    def test_keeps_letter_size_and_one_inch_margins(self, api_client, reports):
        doc = _export(api_client, reports)
        section = doc.sections[0]

        assert section.page_width == 7772400  # Letter
        assert section.left_margin == 914400  # 1"

    def test_carries_no_leftovers_from_the_template_examples(self, api_client, reports):
        """La plantilla trae reportes de muestra que no deben viajar."""
        assert "Las Charcas" not in _todo_el_texto(_export(api_client, reports))

    def test_the_body_is_not_justified(self, api_client, reports):
        """Justificar abre ríos de espacio en cuanto cambia el ancho de columna
        entre Word y Pages."""
        doc = _export(api_client, reports)

        for parrafo in doc.paragraphs:
            assert 'w:val="both"' not in parrafo._p.xml


class TestCover:
    def test_opens_with_the_title_and_what_the_document_covers(self, api_client, reports):
        doc = _export(api_client, reports)

        assert doc.paragraphs[0].text == "Reportes de prensa"
        assert doc.paragraphs[0].style.name == "ODIN Title"
        assert doc.paragraphs[1].style.name == "ODIN Subtitle"
        assert doc.paragraphs[1].text == (
            "ODIN · 21–28 de agosto de 2026 · 2 reportes · Listín Diario, Diario Libre"
        )

    def test_a_single_report_is_not_called_reportes(self, api_client, reports):
        doc = _export(api_client, reports[:1])

        assert "1 reporte ·" in doc.paragraphs[1].text


class TestFicha:
    def test_one_card_per_report(self, api_client, reports):
        doc = _export(api_client, reports)

        assert len(doc.tables) == 2

    def test_the_band_numbers_the_report_and_states_its_sentiment(self, api_client, reports):
        doc = _export(api_client, reports)

        assert _rows(doc.tables[0])[0] == ["REPORTE 01", "NEGATIVO"]
        assert _rows(doc.tables[1])[0] == ["REPORTE 02", "POSITIVO"]

    def test_the_headline_spans_the_whole_card(self, api_client, reports):
        doc = _export(api_client, reports)

        assert _rows(doc.tables[0])[1] == ["Titular del reporte 1"]

    def test_metadata_follows_the_fixed_order(self, api_client, reports):
        """Se lee en Z: quien revisa treinta fichas busca el medio siempre en el
        mismo sitio."""
        doc = _export(api_client, reports)
        casillas = [_field(c) for row in doc.tables[0].rows[2:5] for c in row.cells]

        assert casillas == [
            ("MEDIO", "Listín Diario"),
            ("SECCIÓN", "Provincias"),
            ("PUBLICADO", "21/08/2026"),
            ("TEMA", "tema suelto 1"),
            ("DOCUMENTALISTA", "Juan Pérez"),
            ("ANALIZADO", "28/08/2026"),
        ]

    def test_medio_shows_the_readable_name_not_the_slug(self, api_client, reports):
        """El documento lo lee el cliente: "Listín Diario", no `listin_diario`."""
        texto = _todo_el_texto(_export(api_client, reports))

        assert "Listín Diario" in texto
        assert "listin_diario" not in texto

    def test_lists_the_entities_with_their_count(self, api_client, reports):
        doc = _export(api_client, reports)
        etiqueta, valor = _field(doc.tables[0].rows[5].cells[0])

        assert etiqueta == "ENTIDADES MENCIONADAS · 2"
        assert valor == "Coraasan · Ramón Peña"

    def test_closes_with_the_source_url(self, api_client, reports):
        doc = _export(api_client, reports)

        assert _rows(doc.tables[0])[-1] == ["FUENTE   https://listindiario.com/n1"]

    def test_no_row_can_break_across_pages(self, api_client, reports):
        """Una ficha partida al pie de página es el defecto que el cuadro existe
        para evitar."""
        doc = _export(api_client, reports)

        for cuadro in doc.tables:
            for row in cuadro.rows:
                assert row._tr.find(qn("w:trPr")).find(qn("w:cantSplit")) is not None

    def test_the_band_is_dark_with_white_text(self, api_client, reports):
        """En Pages, un texto sin color explícito sobre un relleno oscuro se
        vuelve negro sobre negro."""
        doc = _export(api_client, reports)
        celda = doc.tables[0].rows[0].cells[0]

        shd = celda._tc.find(qn("w:tcPr")).find(qn("w:shd"))
        assert shd.get(qn("w:fill")) == "22262A"
        assert str(celda.paragraphs[0].runs[0].font.color.rgb) == "FFFFFF"


class TestAReportWithNothingToShow:
    def test_skips_the_entities_row_instead_of_leaving_it_empty(self, api_client, reports):
        doc = _export(api_client, reports)

        assert len(doc.tables[1].rows) == len(doc.tables[0].rows) - 1
        assert "ENTIDADES" not in _rows(doc.tables[1])[-2][0]

    def test_says_automatico_when_no_one_reviewed_it(self, api_client, reports):
        doc = _export(api_client, reports)
        documentalista, analizado = (_field(c) for c in doc.tables[1].rows[4].cells)

        assert documentalista == ("DOCUMENTALISTA", "Automático")
        assert analizado == ("ANALIZADO", "—")

    def test_never_prints_none_or_an_empty_cell(self, api_client, reports):
        doc = _export(api_client, reports)

        for row in doc.tables[1].rows:
            for cell in row.cells:
                assert cell.text.strip()
                assert "None" not in cell.text


class TestMultipleReports:
    def test_each_report_after_the_first_starts_on_a_new_page(self, api_client, reports):
        """Salto explícito y no `page_break_before`: el cuadro es una tabla, y
        una tabla no admite esa propiedad."""
        doc = _export(api_client, reports)
        saltos = [p for p in doc.paragraphs if 'w:type="page"' in p._p.xml]

        assert len(saltos) == 1


class TestBodyIsClean:
    def _body(self, doc) -> list[str]:
        return [p.text for p in doc.paragraphs if p.style.name == "ODIN Body"]

    def test_drops_the_topic_and_the_repeated_headline(self, api_client, reports):
        cuerpo = self._body(_export(api_client, reports))

        assert "tema suelto 1" not in cuerpo
        assert "Titular del reporte 1" not in cuerpo

    def test_drops_the_section_kicker(self, api_client, reports):
        assert "Alegada negligencia" not in self._body(_export(api_client, reports))

    def test_keeps_the_real_paragraphs(self, api_client, reports):
        cuerpo = self._body(_export(api_client, reports))

        assert "Primer párrafo real de la nota, con su punto final." in cuerpo
        assert "Segundo párrafo real de la nota, también terminado." in cuerpo


# Word valida el orden de los hijos de cada elemento de propiedades: uno fuera
# de sitio abre el documento con una advertencia de contenido ilegible, que es
# justo lo que el cliente no puede ver al abrir su reporte. El export escribe
# OOXML a mano (rellenos, bordes, `cantSplit`, tracking) porque python-docx no
# los expone, así que el orden no lo garantiza nadie más que esta prueba.
_SECUENCIAS = {
    "pPr": (
        "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
        "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
        "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE",
        "autoSpaceDN", "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
        "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc", "textDirection",
        "textAlignment", "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
        "sectPr", "pPrChange",
    ),
    "rPr": (
        "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
        "dstrike", "outline", "shadow", "emboss", "imprint", "noProof", "snapToGrid",
        "vanish", "webHidden", "color", "spacing", "w", "kern", "position", "sz",
        "szCs", "highlight", "u", "effect", "bdr", "shd", "fitText", "vertAlign",
        "rtl", "cs", "em", "lang", "eastAsianLayout", "specVanish", "oMath",
    ),
    "tblPr": (
        "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
        "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
        "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription",
        "tblPrChange",
    ),
    "trPr": (
        "cnfStyle", "divId", "gridBefore", "gridAfter", "wBefore", "wAfter", "cantSplit",
        "trHeight", "tblHeader", "tblCellSpacing", "jc", "hidden", "ins", "del",
        "trPrChange",
    ),
    "tcPr": (
        "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd", "noWrap",
        "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
    ),
    "tblBorders": ("top", "left", "bottom", "right", "insideH", "insideV"),
    "tcBorders": ("top", "left", "bottom", "right", "insideH", "insideV", "tl2br", "tr2bl"),
    "tblCellMar": ("top", "left", "bottom", "right"),
    "tcMar": ("top", "left", "bottom", "right"),
}

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class TestOpensWithoutWarnings:
    def test_every_properties_element_follows_the_schema_order(self, api_client, reports):
        doc = _export(api_client, reports)

        for contenedor, secuencia in _SECUENCIAS.items():
            for elemento in doc.element.body.iter(f"{_W}{contenedor}"):
                hijos = [c.tag.removeprefix(_W) for c in elemento]
                posiciones = [secuencia.index(h) for h in hijos if h in secuencia]

                assert posiciones == sorted(posiciones), f"w:{contenedor} desordenado: {hijos}"

    def test_bold_is_never_switched_off_explicitly(self, api_client, reports):
        """`<w:b w:val="0"/>` es peor que no poner nada: el lector de .docx de
        macOS lo lee como negrita ENCENDIDA y saca el documento entero en bold.
        """
        doc = _export(api_client, reports)

        for b in doc.element.body.iter(f"{_W}b"):
            assert b.get(f"{_W}val") in (None, "1", "true")

    def test_the_typography_does_not_depend_on_named_styles(self, api_client, reports):
        """Ese mismo lector ignora los estilos de párrafo propios: sin formato
        directo, la portada sale en Times 12."""
        doc = _export(api_client, reports)
        portada, subtitulo = doc.paragraphs[0], doc.paragraphs[1]

        for parrafo in (portada, subtitulo):
            fuente = parrafo.runs[0].font
            assert fuente.name and fuente.size and fuente.color.rgb

    def test_the_card_is_never_the_last_thing_in_the_document(self, api_client, reports):
        """Una tabla pegada al final del cuerpo hace que Word abra con
        advertencia; siempre queda un párrafo detrás."""
        doc = _export(api_client, reports)
        bloques = [c.tag.removeprefix(_W) for c in doc.element.body]

        assert bloques[-1] == "sectPr"
        assert bloques[-2] == "p"
