"""Pruebas de la exportación a Word.

El caso del cliente: el admin filtra los reportes de un documentalista, selecciona
los que quiere y los baja en un documento.
"""
from __future__ import annotations

import io
from datetime import date, datetime

import pytest
from docx import Document

from odin.core import auth
from odin.core.auth import create_token
from odin.db.models import Article, User


def _headers(username: str = "jperez"):
    token, _ = create_token(username)
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def reports(sqlite_sessionmaker):
    session = sqlite_sessionmaker()
    session.add(
        User(
            username="jperez",
            username_key="jperez",
            display_name="Juan Pérez",
            password_hash=auth.hash_password("x", iterations=1000),
            role="documentalista",
        )
    )
    session.commit()
    juan = session.query(User).filter_by(username="jperez").one().id

    ids = []
    for n in (1, 2):
        article = Article(
            source="listin_diario",
            url=f"https://listindiario.com/e{n}",
            title=f"Reporte número {n}",
            body="Cuerpo de la nota.",
            main_topic="agua potable",
            overall_sentiment="NEG",
            published_at=datetime(2026, 8, n),
            documentalist_id=juan,
            analyzed_on=date(2026, 8, 20),
        )
        session.add(article)
        session.commit()
        ids.append(article.id)
    session.close()
    return ids


def _read(resp) -> Document:
    return Document(io.BytesIO(resp.content))


def _text(resp) -> str:
    """Todo el texto del documento.

    Los datos del reporte viven en el cuadro de ficha, que es una tabla, así que
    mirar solo `doc.paragraphs` dejaría fuera justo lo que estas pruebas
    verifican.
    """
    doc = _read(resp)
    parrafos = [p.text for p in doc.paragraphs]
    celdas = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    return "\n".join(parrafos + celdas)


class TestExport:
    def test_returns_a_word_document(self, api_client, reports):
        resp = api_client.post(
            "/api/articles/export", json={"article_ids": reports}, headers=_headers()
        )

        assert resp.status_code == 200
        assert "wordprocessingml" in resp.headers["content-type"]
        assert ".docx" in resp.headers["content-disposition"]

    def test_includes_only_the_selected_reports(self, api_client, reports):
        resp = api_client.post(
            "/api/articles/export", json={"article_ids": [reports[0]]}, headers=_headers()
        )

        text = _text(resp)
        assert "Reporte número 1" in text
        assert "Reporte número 2" not in text

    def test_carries_the_documentalist_and_the_analysis_date(self, api_client, reports):
        """Es el dato que hace útil el documento cuando se exporta el trabajo
        de una persona."""
        resp = api_client.post(
            "/api/articles/export", json={"article_ids": reports}, headers=_headers()
        )

        text = _text(resp)
        assert "Juan Pérez" in text
        assert "20/08/2026" in text

    def test_shows_the_date_without_a_time(self, api_client, reports):
        resp = api_client.post(
            "/api/articles/export", json={"article_ids": reports}, headers=_headers()
        )

        text = _text(resp)
        assert "00:00" not in text

    def test_rejects_an_empty_selection(self, api_client, reports):
        resp = api_client.post(
            "/api/articles/export", json={"article_ids": []}, headers=_headers()
        )

        assert resp.status_code == 422

    def test_ignores_ids_that_do_not_exist(self, api_client, reports):
        """Un id borrado entre que se listó y se exportó no puede tumbar la
        descarga entera."""
        resp = api_client.post(
            "/api/articles/export",
            json={"article_ids": [reports[0], 999999]},
            headers=_headers(),
        )

        assert resp.status_code == 200
        text = _text(resp)
        assert "Reporte número 1" in text

    def test_fails_when_nothing_selected_exists(self, api_client, reports):
        resp = api_client.post(
            "/api/articles/export", json={"article_ids": [999999]}, headers=_headers()
        )

        assert resp.status_code == 404

    def test_requires_authentication(self, api_client, reports):
        resp = api_client.post("/api/articles/export", json={"article_ids": reports})

        assert resp.status_code in (401, 403)

    def test_medio_shows_the_readable_name_not_the_slug(self, api_client, reports):
        """El documento lo lee el cliente: "Listín Diario", no `listin_diario`."""
        resp = api_client.post(
            "/api/articles/export", json={"article_ids": reports}, headers=_headers()
        )

        text = _text(resp)
        assert "Listín Diario" in text
        assert "listin_diario" not in text
